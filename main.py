from flask import Flask, request
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email.mime.application import MIMEApplication
from email import encoders
import threading
import os
import time
import uuid
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import subprocess



def send_email_with_attachment(name,sender_email, sender_password, receiver_email, subject, body, file_path):
    # Setup the MIME
    message = MIMEMultipart()
    message['From'] = sender_email
    message['To'] = receiver_email
    message['Subject'] = subject

    # Attach body
    message.attach(MIMEText(body, 'plain'))

    with open(file_path, 'rb') as file:
        attach = MIMEApplication(file.read(),_subtype="pdf")
        attach.add_header('Content-Disposition','attachment',filename=str(name)+'.pdf')
    message.attach(attach)

    # Connect to the SMTP server and send the email
    with smtplib.SMTP('smtp.gmail.com', 587) as server:
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, receiver_email, message.as_string())

app = Flask(__name__)

@app.route('/webhook', methods=['POST'])
def handle_webhook():
    # Handle the incoming GET request data
    data = request.json 
    
    # Process the received data (you can perform actions based on this data)
    # For now, just print the received data
    print("Received form data:", data)
    form_data = data['data']
    # Add your handling logic here
    doc = Document('template.docx')
    # styled_font = {
    # 'bold': False,
    # 'italic': False,
    # 'underline': False,
    # 'color': '000000',  # Red color (format: 'AARRGGBB')
    # 'name': 'Arial',
    # 'size': 12
    # }
    find_replace_word(doc, '$NAME$', form_data['姓名'])
    find_replace_word(doc, '$GENDER$', form_data['性別'])
    find_replace_word(doc, '$Y$', form_data['出生年月日'][0:4])
    find_replace_word(doc, '$M$', form_data['出生年月日'][5:7])
    find_replace_word(doc, '$D$', form_data['出生年月日'][8:10])
    find_replace_word(doc, '$ID$', form_data['身分證字號'])
    find_replace_word(doc, '$PHONE$', form_data['聯絡電話'])
    find_replace_word(doc, '$CONTACTADDRESS$', form_data['通訊地址'])
    find_replace_word(doc, '$HOMEADDRESS$', form_data['戶籍地址'])
    find_replace_word(doc, '$EMAIL$', data['email'])
    find_replace_word(doc, '$SCHOOL$', form_data['畢業學校'])
    find_replace_word(doc, '$SUBJECT$', form_data['畢業系所'])
    find_replace_word(doc, '$CAREER$', form_data['現職'])
    find_replace_word(doc, '$GY$', form_data['畢業年份'])
    find_replace_word(doc, '$LINE$', form_data['Line ID'])
    find_replace_word(doc, '$SID$', form_data['在學期間的學號'])
    name= form_data['姓名']
    receiver_email= data['email']
    email_subject = '【陽明交通大學校友會】校友資料確認信'
    email_body='敬愛的' + name + '校友夥伴您好：\n\n我們已收到您送出的會員入會表單，感謝您對陽明交大校友總會的鼎力支持。\n\n附件為包含您個人資料的申請表檔案，請您下載並列印該份檔案、確認資料無誤（或以紅筆修改有誤的資料）並簽章後，連同申請表上所列的附件，於113年1月31日前以掛號寄回本會（郵寄資料載於申請表第二頁）。我們會在收到書面申請表後再次以電子郵件通知您。\n\n如您遇到任何問題，歡迎隨時與我們聯繫：\n校友總會籌備會電子信箱｜nycualumni@gmail.com\n校友總會籌備會粉絲專頁｜https://www.facebook.com/nycualumni\n\n您的加入與支持，將成為校友總會最堅實的後盾。我們因為對母校的情感而齊聚、為了回饋母校而共同努力，相信我們今日的連結，都將成為同行於社會上的力量。\n\n國立陽明交通大學校友總會籌備會\n主任委員　梁家語' 
    uuid4 = str(uuid.uuid4())
    doc.save(uuid4+name + '.docx')
    fiiename= uuid4+name + '.docx'
    file_path = './nycuaa/auto_mail/'+name + '.docx'
    task_thread = threading.Thread(target=send_request_and_process, args=(name, receiver_email, email_subject, email_body, fiiename))
    task_thread.start()
    return "Received the data successfully!"

def send_request_and_process(name, receiver_email, email_subject, email_body, fiiename):
    while not os.path.exists(fiiename):
        print('wait for ' + fiiename + '...')
        time.sleep(1)
    time.sleep(1)
    pdfname= fiiename[:-4] + 'pdf'
    subprocess.run(['doc2pdf', fiiename])
    # convert_docx_to_pdf(fiiename, pdfname)
    # time.sleep(5)
    print('send email to ' + receiver_email + '...')
    send_email_with_attachment(name,'nycualumni@gmail.com', 'vffgwnqaldkjaebv', receiver_email, email_subject, email_body, pdfname) 
    print('send email to ' + receiver_email + ' successfully!')

def find_replace_word(doc, old_word, new_word):
    for paragraph in doc.paragraphs:
        if old_word in paragraph.text:
            paragraph.text = paragraph.text.replace(old_word, new_word)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_word in cell.text:
                    cell.text = cell.text.replace(old_word, new_word)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)